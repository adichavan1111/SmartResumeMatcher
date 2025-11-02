import streamlit as st
import pandas as pd
import re
import smtplib
from email.mime.text import MIMEText
from sentence_transformers import SentenceTransformer, util
from docx import Document
import fitz  # PyMuPDF for PDFs
import tempfile
import os
import pythoncom
from docx2pdf import convert
import pypandoc  # fallback

# -------------------------------------------------
# 🔹 Helper Functions
# -------------------------------------------------
def get_text_from_pdf(pdf_path):
    """Extract text from a PDF file path."""
    text = ""
    with fitz.open(pdf_path) as pdf:
        for page in pdf:
            text += page.get_text()
    return text


def extract_mobile_number(text):
    match = re.search(r'(\+?\d{1,3}[-.\s]?)?(\d{10}|\d{3}[-.\s]\d{3}[-.\s]\d{4})', text)
    return match.group(0) if match else "Not Found"


def extract_email(text):
    match = re.search(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else "Not Found"


def send_mail_mailtrap(to_email, subject, body):
    """Send email via Mailtrap."""
    sender = "recruiter@yourcompany.com"
    msg = MIMEText(body, "plain")
    msg["Subject"] = subject
    msg["From"] = sender
    msg["To"] = to_email

    SMTP_HOST = "sandbox.smtp.mailtrap.io"
    SMTP_PORT = 587
    SMTP_USER = "YOUR_MAILTRAP_USERNAME"
    SMTP_PASS = "YOUR_MAILTRAP_PASSWORD"

    try:
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT, timeout=20) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.send_message(msg)
        st.success(f"✅ Email sent successfully to {to_email}")
    except Exception as e:
        st.error(f"❌ Failed to send to {to_email}: {e}")


def safe_convert_docx_to_pdf(docx_filename, pdf_filename):
    """Safely convert DOCX → PDF using docx2pdf with pythoncom, fallback to pypandoc."""
    try:
        pythoncom.CoInitialize()
        convert(docx_filename, pdf_filename)
        pythoncom.CoUninitialize()
        st.success(f"✅ Converted {os.path.basename(docx_filename)} → PDF using Microsoft Word")
        return pdf_filename
    except Exception as e:
        st.warning(f"⚠️ docx2pdf failed: {e}. Trying fallback (pypandoc)...")
        try:
            pypandoc.convert_file(docx_filename, 'pdf', outputfile=pdf_filename, extra_args=['--standalone'])
            st.success("✅ Converted using pypandoc fallback")
            return pdf_filename
        except Exception as e2:
            st.error(f"❌ Both conversions failed: {e2}")
            return None


def convert_docx_to_pdf_bytes(uploaded_file):
    """Convert an uploaded DOCX file into PDF bytes using temporary files."""
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(uploaded_file.read())
        temp_docx_path = temp_docx.name

    pdf_output = temp_docx_path.replace(".docx", ".pdf")
    pdf_path = safe_convert_docx_to_pdf(temp_docx_path, pdf_output)
    if pdf_path and os.path.exists(pdf_path):
        with open(pdf_path, "rb") as pdf_file:
            pdf_bytes = pdf_file.read()
        return pdf_path, pdf_bytes
    return None, None


# -------------------------------------------------
# 🔹 NLP Utilities
# -------------------------------------------------
def preprocess_text(text):
    """Clean and extract skill-related keywords from text."""
    text = re.sub(r'[^a-zA-Z0-9\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    keywords = re.findall(r'\b(Java|Python|AWS|SQL|Testing|Machine Learning|AI|Automation|DevOps|Selenium|C\+\+|Cloud)\b', text, flags=re.IGNORECASE)
    return " ".join(keywords) if keywords else text


@st.cache_resource
def load_model():
    # ⚡ Better model for resume-job matching
    return SentenceTransformer("all-mpnet-base-v2")

model = load_model()

def get_embedding(text, model):
    return model.encode(text, convert_to_tensor=True, normalize_embeddings=True)


# -------------------------------------------------
# 🔹 Streamlit UI Setup
# -------------------------------------------------
st.set_page_config(page_title="AI Resume Matcher", layout="wide")
st.title("🤖 AI-Powered Resume Matcher")

st.markdown("""
### 💼 Business Description  
This AI-powered Resume Matcher uses **advanced NLP (SentenceTransformer)** to go **beyond keyword search** — it understands context, skills, and job relevance to match resumes with job descriptions intelligently.  
It employs **semantic similarity**, extracting meaningful insights from each document to help recruiters shortlist top candidates faster.
""")


# -------------------------------------------------
# 🔹 Upload Section
# -------------------------------------------------
job_description = st.text_area("📄 Paste Job Description", height=200)
uploaded_resumes = st.file_uploader("📤 Upload Resumes (PDF or DOCX)", accept_multiple_files=True)


# -------------------------------------------------
# 🔹 Process Resumes
# -------------------------------------------------
if st.button("🚀 Match Resumes with Job Description"):
    if not job_description or not uploaded_resumes:
        st.warning("Please upload resumes and enter a job description.")
    else:
        st.info("⏳ Processing resumes... Please wait...")

        job_text = preprocess_text(job_description)
        job_embedding = get_embedding(job_text, model)
        results = []

        for resume in uploaded_resumes:
            if resume.name.endswith(".docx"):
                st.write(f"📄 Converting {resume.name} → PDF...")
                pdf_path, _ = convert_docx_to_pdf_bytes(resume)
                if pdf_path:
                    text = get_text_from_pdf(pdf_path)
                else:
                    text = ""
            elif resume.name.endswith(".pdf"):
                with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_pdf:
                    temp_pdf.write(resume.read())
                    text = get_text_from_pdf(temp_pdf.name)
            else:
                st.warning(f"⚠️ Unsupported file type: {resume.name}")
                continue

            if text.strip():
                clean_text = preprocess_text(text)
                resume_embedding = get_embedding(clean_text, model)
                score = util.cos_sim(job_embedding, resume_embedding)[0][0].item() * 100

                mobile = extract_mobile_number(text)
                email = extract_email(text)
                results.append({
                    "Resume Name": resume.name,
                    "Match %": round(score, 2),
                    "Mobile Number": mobile,
                    "Email": email
                })
            else:
                results.append({
                    "Resume Name": resume.name,
                    "Match %": 0,
                    "Mobile Number": "-",
                    "Email": "-"
                })

        # -------------------------------------------------
        # 🔹 Prepare DataFrame
        # -------------------------------------------------
        df = pd.DataFrame(results).sort_values(by="Match %", ascending=False)
        threshold = max(25, df["Match %"].mean() * 0.8)
        df["Status"] = df["Match %"].apply(lambda x: "✅ Matched" if x >= threshold else "❌ Not Matched")

        df.insert(0, "S.No", range(1, len(df) + 1))
        st.session_state["match_df"] = df

        st.subheader("📊 Match Results")
        st.dataframe(df, use_container_width=True)

        # -------------------------------------------------
        # 🔹 Export Report (DOCX + PDF)
        # -------------------------------------------------
        docx_filename = "resume_match_report.docx"
        pdf_filename = "resume_match_report.pdf"

        doc = Document()
        doc.add_heading("Resume Match Report", level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(df.columns):
            hdr_cells[i].text = col_name

        # Data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        doc.save(docx_filename)
        safe_convert_docx_to_pdf(docx_filename, pdf_filename)

        # -------------------------------------------------
        # 🔹 Download Buttons
        # -------------------------------------------------
        if os.path.exists(pdf_filename):
            with open(pdf_filename, "rb") as f:
                st.download_button("⬇️ Download PDF Report", f, file_name=pdf_filename, mime="application/pdf")

        csv = df.to_csv(index=False).encode("utf-8")
        st.download_button(
            label="📥 Download Match Report (CSV)",
            data=csv,
            file_name="resume_match_report.csv",
            mime="text/csv",
        )















